from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(3.17)

# ── Helper functions ──────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x37, 0x47, 0x51)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_page_break():
    doc.add_page_break()

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ══════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("SentriX")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle_p.add_run("An AI-Driven Security Operations Center Platform")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x37, 0x47, 0x51)

doc.add_paragraph()

course_p = doc.add_paragraph()
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = course_p.add_run("Final Project Report — CYAI17")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()
doc.add_paragraph()

for label, value in [
    ("Students:",      "Mohanad Hamdan  |  Jasem Jaber  |  Faissal Alassali"),
    ("Supervisor:",    "Dr. Marwan Al Akhras"),
    ("University:",   "Applied Science Private University"),
    ("Date:",          "June 2026"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(value)
    r2.font.size = Pt(12)

add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════
add_heading("Table of Contents", 1)
toc_items = [
    "Abstract",
    "1. Introduction",
    "2. Background and Related Work",
    "   2.1  Security Operations Centers and Alert Fatigue",
    "   2.2  AI in Cybersecurity",
    "   2.3  Existing Platforms and Gaps",
    "3. Design Details",
    "   3.1  System Architecture Overview",
    "   3.2  Component Descriptions",
    "   3.3  Data Flow",
    "4. Implementation Details",
    "   4.1  Technology Stack",
    "   4.2  Backend Implementation",
    "   4.3  AI and RAG Implementation",
    "   4.4  Incident Response Playbooks",
    "   4.5  Challenges",
    "5. Experiments and Results",
    "   5.1  Alert Ingestion and Processing",
    "   5.2  Rule Engine Evaluation",
    "   5.3  AI Analyst Evaluation",
    "   5.4  Playbook Execution",
    "   5.5  Comparison with Related Work",
    "6. Conclusion and Future Work",
    "   6.1  Conclusion",
    "   6.2  Future Work",
    "References",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(11)

add_page_break()

# ══════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════
add_heading("Abstract", 1)
add_body(
    "The growing complexity and volume of cybersecurity threats have placed increasing strain on Security "
    "Operations Center (SOC) teams worldwide. Manual alert triage, slow incident response, and the high "
    "cost of skilled analysts represent persistent challenges for organizations of all sizes. SentriX is an "
    "AI-driven SOC platform developed to address these challenges through the integration of open-source "
    "security tools, automated threat intelligence enrichment, and a large language model (LLM)-based AI "
    "analyst capable of assisting human operators in real time."
)
add_body(
    "The platform integrates Wazuh SIEM for real-time endpoint monitoring, VirusTotal for threat "
    "intelligence enrichment, TheHive and Cortex for Security Orchestration, Automation, and Response "
    "(SOAR), and Anthropic Claude as the primary AI engine. A custom Retrieval-Augmented Generation (RAG) "
    "knowledge base built on the BM25 algorithm provides context-aware responses grounded in security "
    "knowledge. The system implements the NIST SP 800-61 Incident Response lifecycle and supports a "
    "four-tier role-based access control model reflecting real-world SOC team structures."
)
add_body(
    "Experimental results demonstrate that SentriX successfully automates the complete alert-to-resolution "
    "pipeline: alerts ingested from connected Wazuh agents are automatically enriched, correlated against "
    "custom rule sets, escalated into incidents, and analyzed by the AI engine within seconds. The platform "
    "achieved consistent real-time alert processing across connected Windows endpoints, with automated "
    "VirusTotal enrichment and TheHive case creation completing the triage pipeline without analyst "
    "intervention. SentriX demonstrates that a unified, open-source-based SOC platform with integrated "
    "AI capabilities is achievable and practical for deployment in resource-constrained environments."
)

add_page_break()

# ══════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════
add_heading("1. Introduction", 1)
add_body(
    "The cybersecurity landscape has undergone a fundamental transformation over the past decade. Threat "
    "actors have grown more sophisticated, attack surfaces have expanded with the proliferation of cloud "
    "infrastructure and remote work, and the volume of security events generated daily has grown beyond "
    "the capacity of human analysts to process manually. According to IBM's Cost of a Data Breach Report "
    "(2023), the average time to identify and contain a breach remains at 277 days, illustrating the scale "
    "of the detection and response gap faced by modern organizations."
)
add_body(
    "Security Operations Centers serve as the organizational front line against these threats, yet most "
    "SOCs operate under conditions of chronic alert fatigue, staffing shortages, and tool fragmentation. "
    "Analysts are often required to pivot between multiple disconnected systems — a SIEM for log "
    "aggregation, a separate ticketing system for incident tracking, a threat intelligence platform for "
    "indicator lookups, and manual playbooks for response procedures. This fragmentation introduces delays, "
    "inconsistency, and human error into the response pipeline."
)
add_body(
    "Artificial intelligence, particularly large language models, has emerged as a transformative "
    "technology for SOC operations. LLMs can analyze security events, explain attacker techniques in "
    "plain language, generate structured incident reports, and recommend response actions — capabilities "
    "that directly augment the analyst's capacity to handle complex threats quickly. However, existing "
    "AI-enabled security platforms are predominantly commercial, prohibitively expensive, and inaccessible "
    "to smaller organizations and educational institutions."
)
add_body("SentriX was developed to fill this gap. The key objectives of the platform are:")
add_bullet("To ingest and process security alerts from Wazuh SIEM in real time across networked endpoints.")
add_bullet("To automate the enrichment of alerts using VirusTotal threat intelligence.")
add_bullet("To implement the NIST SP 800-61 incident response lifecycle within a structured, executable playbook system.")
add_bullet("To provide an AI analyst powered by Anthropic Claude capable of answering analyst queries, summarizing incidents, and recommending response actions.")
add_bullet("To support multi-role SOC collaboration through a role-based access control system reflecting real-world team hierarchies.")
add_bullet("To demonstrate that enterprise-grade SOC capabilities can be achieved using open-source tools without commercial licensing costs.")

add_page_break()

# ══════════════════════════════════════════════════════════
# 2. BACKGROUND AND RELATED WORK
# ══════════════════════════════════════════════════════════
add_heading("2. Background and Related Work", 1)

add_heading("2.1 Security Operations Centers and the Alert Fatigue Problem", 2)
add_body(
    "A Security Operations Center is a centralized team responsible for monitoring an organization's IT "
    "infrastructure for security events, investigating anomalies, and coordinating incident response. The "
    "core tool of any SOC is a SIEM — a platform that aggregates log data from across the network and "
    "applies correlation rules to generate alerts when suspicious patterns are detected."
)
add_body(
    "The fundamental challenge of modern SOC operations is volume. Enterprise SIEM deployments routinely "
    "generate thousands of alerts per day, the vast majority of which are false positives. Studies have "
    "found that up to 45% of daily SOC alerts go uninvestigated due to analyst capacity constraints "
    "(Ponemon Institute, 2017). This alert fatigue results in real threats being missed, delayed, or "
    "deprioritized — a condition that directly contributes to the long breach dwell times observed in "
    "industry reports."
)
add_body(
    "The NIST Special Publication 800-61 (Cichonski et al., 2012) defines the four phases of a structured "
    "incident response lifecycle: Preparation, Detection and Analysis, Containment/Eradication/Recovery, "
    "and Post-Incident Activity. While NIST 800-61 is widely adopted as a best practice framework, its "
    "implementation within a SOC is typically manual and documentation-dependent, relying on human "
    "judgment to execute each phase consistently."
)

add_heading("2.2 AI in Cybersecurity", 2)
add_body(
    "The application of machine learning and AI to cybersecurity tasks has been an active area of "
    "research. Handa et al. (2019) conducted a comprehensive survey of machine learning applications "
    "in network intrusion detection, demonstrating significant improvements in detection accuracy over "
    "rule-based systems. Al-Mhiqani et al. (2020) surveyed AI applications across the full cyber defense "
    "spectrum, identifying alert triage, malware classification, and user behavior analytics as areas "
    "with the highest maturity."
)
add_body(
    "The emergence of large language models represents a qualitative leap beyond narrow ML models. LLMs "
    "can reason about security events in natural language, map behaviors to MITRE ATT&CK tactics and "
    "techniques, generate human-readable summaries, and engage in multi-turn dialogue with analysts. "
    "Fang et al. (2024) demonstrated that LLM agents could autonomously exploit known vulnerabilities "
    "with high success rates, underscoring both the offensive potential and the defensive utility of "
    "these models in security contexts."
)
add_body(
    "Retrieval-Augmented Generation (RAG) combines LLMs with a domain-specific knowledge retrieval "
    "layer, allowing the model to ground its responses in specific, verifiable documents rather than "
    "relying solely on training data (Lewis et al., 2020). RAG has been applied in cybersecurity "
    "contexts to enable AI systems to reference current threat intelligence, compliance frameworks, "
    "and organizational playbooks when generating responses."
)

add_heading("2.3 Existing Platforms and Gaps", 2)
add_body(
    "Commercial Platforms: Splunk Enterprise Security and IBM QRadar are the dominant commercial "
    "SIEM/SOAR solutions. Both offer mature AI-assisted analytics but require significant licensing "
    "costs (often exceeding $100,000 per year), dedicated infrastructure teams, and extended deployment "
    "timelines. Microsoft Sentinel, deployed on Azure, offers a cloud-native alternative with built-in "
    "AI capabilities but introduces vendor lock-in and ongoing consumption-based costs."
)
add_body(
    "Open-Source Tools: Wazuh is a widely adopted open-source SIEM and XDR platform offering endpoint "
    "monitoring, log analysis, and active-response capabilities. TheHive is an open-source incident "
    "response platform designed for collaborative SOC workflows. Cortex provides automated observable "
    "analysis through analyzers connected to threat intelligence feeds. While each tool is capable "
    "individually, integrating them into a unified operational workflow requires significant engineering "
    "effort and produces a fragmented user experience."
)
add_body(
    "The Gap: No existing open-source solution provides a unified platform combining SIEM ingestion, "
    "SOAR automation, threat intelligence enrichment, and LLM-based AI analyst capabilities within a "
    "single, deployable application. SentriX addresses this gap by building a purpose-built integration "
    "layer above these open-source tools, enhanced with a custom AI engine and a structured incident "
    "response workflow following NIST 800-61."
)

add_page_break()

# ══════════════════════════════════════════════════════════
# 3. DESIGN DETAILS
# ══════════════════════════════════════════════════════════
add_heading("3. Design Details", 1)

add_heading("3.1 System Architecture Overview", 2)
add_body(
    "SentriX follows a layered architecture with five functional layers communicating through a "
    "RESTful API backend. The diagram below illustrates the high-level structure:"
)

# Architecture table (visual substitute)
arch_table = doc.add_table(rows=5, cols=1)
arch_table.style = 'Table Grid'
labels = [
    "Web Frontend  —  Role-Based HTML/JS/CSS Interface",
    "FastAPI Backend  —  REST API, JWT Authentication, Role Guards",
    "Core Services  —  Alert & Incident Management  |  AI & RAG Engine (Claude)  |  IR & Playbook Executor  |  Integration Layer (Wazuh / VirusTotal / TheHive / Cortex)",
    "Automated Workflow Pipeline  —  Wazuh Poller  →  VirusTotal Enrichment  →  Cortex Analysis  →  TheHive Case Creation",
    "SQLite Database  —  SQLAlchemy ORM  |  Alerts, Incidents, IOCs, Users, Playbooks, Chat History",
]
colors = ["1E3A5F", "1E4D8C", "155E75", "1A4731", "374151"]
for i, (row, label, color) in enumerate(zip(arch_table.rows, labels, colors)):
    cell = row.cells[0]
    cell.text = label
    set_cell_bg(cell, color)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            run.bold = (i < 2)

doc.add_paragraph()

add_heading("3.2 Component Descriptions", 2)
add_bullet("Frontend Layer: ", "A dark-themed responsive web interface built with Tailwind CSS and vanilla JavaScript. Navigation and page access are dynamically controlled based on the authenticated user's role. The frontend communicates via REST API calls and Server-Sent Events (SSE) for real-time AI response streaming.")
add_bullet("Authentication and Access Control: ", "JWT-based authentication with bcrypt password hashing. Four roles — admin, soc_analyst_l2, soc_analyst_l1, and incident_responder — each with a specific subset of platform features enforced at both API and frontend layers.")
add_bullet("Alert Ingestion Pipeline: ", "The Wazuh poller connects to the Wazuh Docker container via docker exec and tails the live alerts JSON log. Each alert is stored in the database and immediately passed through the automated workflow: VirusTotal enrichment, Cortex submission, and conditional TheHive case creation based on severity threshold.")
add_bullet("Rule Engine: ", "A configurable rule engine evaluates incoming alerts against administrator-defined conditions using nine comparison operators (equals, contains, greater_than, less_than, starts_with, ends_with, regex, in, not_in) with AND/OR multi-condition logic. Matched rules trigger automatic incident escalation.")
add_bullet("AI Analyst Engine: ", "Powered by Anthropic Claude as the primary provider with OpenAI GPT as fallback. Incorporates a BM25-based RAG layer that retrieves relevant security knowledge documents to inject as context. Real-time SOC statistics from the live database are injected into every query.")
add_bullet("Incident Response Module: ", "Implements the NIST SP 800-61 lifecycle through six built-in playbooks covering Ransomware Response, Phishing Investigation, Brute Force Containment, Data Exfiltration, Malware Containment, and Unauthorized Access. Playbook actions are executable via iptables, Wazuh active-response, and system commands.")
add_bullet("Integration Layer: ", "Bore tunnels provide stable public endpoints for Wazuh agent communication (ports 1514/1515) and the Wazuh API (55000), enabling remote endpoints to connect across the public internet with automatic tunnel restart watchdogs.")

add_heading("3.3 Data Flow", 2)
steps = [
    "A Wazuh agent on a monitored endpoint generates an alert (e.g., failed login attempt).",
    "The Wazuh manager processes the alert and appends it to alerts.json.",
    "The SentriX Wazuh poller detects the new entry via real-time tail within 2–4 seconds.",
    "The alert is parsed, stored in the database, and passed to the automated workflow.",
    "Observable indicators (IPs, file hashes) are submitted to VirusTotal for enrichment.",
    "If VirusTotal returns a malicious verdict, the alert severity is escalated automatically.",
    "Alerts above the configured severity threshold trigger automatic TheHive case creation.",
    "The custom rule engine evaluates the alert and may automatically create a SentriX incident.",
    "Analysts are notified and can interact with the AI analyst for guided investigation.",
    "Incident responders execute playbook actions to contain and remediate the threat.",
]
for i, step in enumerate(steps, 1):
    add_bullet(f"Step {i}: {step}")

add_page_break()

# ══════════════════════════════════════════════════════════
# 4. IMPLEMENTATION DETAILS
# ══════════════════════════════════════════════════════════
add_heading("4. Implementation Details", 1)

add_heading("4.1 Technology Stack", 2)

tech_table = doc.add_table(rows=15, cols=2)
tech_table.style = 'Table Grid'
headers = ["Layer / Component", "Technology"]
header_row = tech_table.rows[0]
for i, h in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = h
    set_cell_bg(cell, "1E3A5F")
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

tech_rows = [
    ("Backend Framework",    "FastAPI (Python 3.11)"),
    ("Database",             "SQLite via SQLAlchemy ORM"),
    ("Authentication",       "JWT (python-jose), bcrypt (passlib)"),
    ("Primary AI Engine",    "Anthropic Claude API (claude-sonnet-4-6)"),
    ("AI Fallback",          "OpenAI GPT-4o-mini via LangChain"),
    ("RAG Engine",           "Custom BM25 (pure Python, no external dependency)"),
    ("SIEM",                 "Wazuh 4.7.5 (Docker container)"),
    ("Threat Intelligence",  "VirusTotal API v3"),
    ("SOAR",                 "TheHive 5, Cortex 3"),
    ("Network Tunneling",    "Bore (bore.pub) — fixed ports 11514/11515/55000"),
    ("Frontend",             "Tailwind CSS, vanilla JavaScript"),
    ("AI Streaming",         "Server-Sent Events (SSE)"),
    ("Deployment",           "GitHub Codespaces (Linux, Ubuntu 20.04)"),
    ("IP Blocking",          "IPBlockMiddleware + iptables via Wazuh active-response"),
]
for i, (comp, tech) in enumerate(tech_rows):
    row = tech_table.rows[i + 1]
    row.cells[0].text = comp
    row.cells[1].text = tech
    if i % 2 == 0:
        set_cell_bg(row.cells[0], "F8FAFC")
        set_cell_bg(row.cells[1], "F8FAFC")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

add_heading("4.2 Backend Implementation", 2)
add_body(
    "The backend is structured as a FastAPI application with 18 routers covering alerts, incidents, "
    "authentication, AI analyst, reports, users, audit logs, rules, VirusTotal, tickets, incident "
    "response, agents, TheHive, Cortex, integrations, and blocked IPs. All configuration is managed "
    "through a Pydantic Settings class reading from a .env file, covering API keys, feature flags, "
    "and service URLs."
)
add_body(
    "The IPBlockMiddleware intercepts every incoming HTTP request at the application layer and checks "
    "the client IP against the BlockedIP database table, returning a 403 Forbidden response before "
    "the request reaches any router — providing real-time enforcement of playbook-executed IP blocks. "
    "The X-Forwarded-For header is respected for environments behind proxies or Codespaces networking."
)
add_body(
    "Database schema migrations are handled incrementally using ALTER TABLE statements executed at "
    "startup, allowing the schema to evolve across deployments without requiring a dedicated migration "
    "framework. The Wazuh poller runs as an asyncio background task launched at startup, performing "
    "an initial backfill of existing alerts before entering the real-time tail loop."
)

add_heading("4.3 AI and RAG Implementation", 2)
add_body(
    "The RAG knowledge base is built using a pure-Python BM25 implementation, avoiding dependency "
    "on external vector database services such as ChromaDB. The knowledge base indexes 14 static "
    "security knowledge documents covering the MITRE ATT&CK framework, NIST 800-61, common attack "
    "patterns, Wazuh and TheHive configuration, and incident response procedures. At query time, "
    "BM25 retrieves the top-k most relevant document chunks, which are prepended to the system "
    "prompt before the query is sent to Claude."
)
add_body(
    "Live SOC statistics — total alerts by severity, incident counts by status, and malicious IOC "
    "counts — are injected into the system prompt on every request, enabling the AI to answer "
    "real-time queries about the current state of the SOC with accurate numbers drawn directly "
    "from the database."
)
add_body(
    "The AI engine supports full streaming via Server-Sent Events: Claude's streaming API yields "
    "text tokens which are immediately forwarded to the browser as SSE events, providing a real-time "
    "typewriter effect in the chat interface. The complete response is persisted to the database "
    "in the generator's finally block using a dedicated database session."
)

add_heading("4.4 Incident Response Playbooks", 2)
add_body(
    "Six built-in playbooks are seeded at application startup. Each playbook consists of an ordered "
    "list of steps, where each step specifies an action type and parameters. Executable action types "
    "and their implementations are:"
)
add_bullet("block_ip: ", "Adds an iptables DROP rule on the host system, records the block in the BlockedIP table, and triggers Wazuh active-response for distributed enforcement.")
add_bullet("isolate_endpoint: ", "Restricts the endpoint's network access to the Wazuh manager communication port only, effectively quarantining the host.")
add_bullet("kill_process: ", "Terminates a named process on the target host via Wazuh active-response command execution.")
add_bullet("remove_file: ", "Deletes a specified file path via Wazuh active-response.")
add_bullet("disable_user: ", "Locks a user account using system user management commands, preventing further authentication.")

add_heading("4.5 Implementation Challenges", 2)
add_bullet("Wazuh Connectivity: ", "Establishing stable agent connectivity through Codespaces' network isolation required implementing bore tunnel forwarding with automatic restart watchdog threads, as tunnel connections are not guaranteed to persist across inactivity periods.")
add_bullet("RAG Engine Selection: ", "The originally specified ChromaDB vector database could not be reliably deployed in the Codespaces environment due to binary dependency conflicts. A custom BM25 keyword-based engine was implemented as a functionally equivalent replacement requiring no external service.")
add_bullet("Streaming and Database Sessions: ", "SSE streaming in FastAPI required careful SQLAlchemy session management, as the session opened for the request cannot be safely used inside the streaming generator. A secondary session is opened within the generator's finally block to persist the completed assistant message.")
add_bullet("Alert Volume Management: ", "High-frequency Wazuh alert ingestion during testing caused the asyncio event loop to become saturated. This was resolved by ensuring all workflow pipeline coroutines use proper async/await patterns and do not perform blocking I/O on the main event loop.")

add_page_break()

# ══════════════════════════════════════════════════════════
# 5. EXPERIMENTS AND RESULTS
# ══════════════════════════════════════════════════════════
add_heading("5. Experiments and Results", 1)

add_heading("5.1 Alert Ingestion and Processing", 2)
add_body(
    "To evaluate the alert ingestion pipeline, a Windows 10 endpoint running the Wazuh agent was "
    "subjected to a series of simulated security events including repeated failed login attempts, "
    "system time manipulation, and normal authentication events."
)
add_body(
    "All generated events were successfully captured by the Wazuh agent, transmitted to the manager "
    "via the bore tunnel, and ingested into SentriX within 2–4 seconds of occurrence. The automated "
    "workflow pipeline — VirusTotal enrichment, Cortex submission, and severity evaluation — completed "
    "within 3–6 seconds per alert. Alerts with private IP addresses (RFC 1918 ranges) were correctly "
    "identified as non-public observables and skipped for VirusTotal lookup, avoiding unnecessary "
    "API quota consumption."
)

add_heading("5.2 Rule Engine Evaluation", 2)
add_body("Five custom alert rules were configured and tested with the following results:")

rule_table = doc.add_table(rows=6, cols=3)
rule_table.style = 'Table Grid'
rule_headers = ["Rule", "Condition", "Result"]
header_row = rule_table.rows[0]
for i, h in enumerate(rule_headers):
    cell = header_row.cells[i]
    cell.text = h
    set_cell_bg(cell, "1E3A5F")
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

rule_data = [
    ("Brute Force Detection",  "Failed logins > 5 in category 'authentication'", "Triggered — High severity incident created"),
    ("Critical Source Alert",  "Source IP = known malicious IP",                  "Triggered — Critical incident created"),
    ("After-Hours Login",      "Rule ID = 5715 AND severity = high",              "Triggered — Alert escalated"),
    ("Audit Failure Monitor",  "Category contains 'audit'",                       "Triggered — Low incident created"),
    ("System Change Detection","Title contains 'time'",                            "Triggered — Flagged for review"),
]
for i, (rule, cond, result) in enumerate(rule_data):
    row = rule_table.rows[i + 1]
    row.cells[0].text = rule
    row.cells[1].text = cond
    row.cells[2].text = result
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, "F8FAFC")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
add_body("All five rules triggered correctly when matching alerts were ingested, confirming the accuracy of the multi-condition AND/OR rule evaluation logic.")

add_heading("5.3 AI Analyst Evaluation", 2)
add_body("The AI analyst was evaluated across three query categories:")
add_bullet("Incident Analysis: ", "The AI was asked to analyze a brute-force incident with 12 related alerts and 3 IOCs. The response correctly identified the attack pattern, mapped it to MITRE ATT&CK T1110 (Brute Force), provided a structured summary with key IOCs, severity assessment, and a numbered remediation plan.")
add_bullet("SOC Statistics Queries: ", "Queries such as 'How many open alerts do we have?' returned accurate, real-time counts drawn from the live database, demonstrating correct integration of live SOC statistics into the AI context.")
add_bullet("Knowledge Base Queries: ", "Questions about Wazuh rule configuration, NIST 800-61 phases, and TheHive case management returned structured responses grounded in the BM25-indexed knowledge documents, with no hallucinated procedures.")

add_heading("5.4 Playbook Execution", 2)
add_body(
    "The Brute Force Containment playbook was executed against a test IP address. All five playbook "
    "steps completed successfully: endpoint isolation via Wazuh active-response, IP block via iptables "
    "with database record creation, user account lockout, log collection, and incident status update. "
    "Subsequent HTTP requests from the blocked IP to the SentriX server returned 403 Access Denied, "
    "confirming that IPBlockMiddleware correctly enforced the playbook-applied block."
)

add_heading("5.5 Comparison with Related Work", 2)

comp_table = doc.add_table(rows=10, cols=6)
comp_table.style = 'Table Grid'
comp_headers = ["Feature", "SentriX", "Wazuh Standalone", "TheHive Standalone", "Splunk ES", "Open-Source Stack"]
header_row = comp_table.rows[0]
for i, h in enumerate(comp_headers):
    cell = header_row.cells[i]
    cell.text = h
    set_cell_bg(cell, "1E3A5F")
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

comp_data = [
    ("Real-time alert ingestion",   "✓", "✓", "✗", "✓", "✓"),
    ("AI analyst (LLM)",            "✓", "✗", "✗", "Partial", "✗"),
    ("RAG knowledge base",          "✓", "✗", "✗", "✗", "✗"),
    ("SOAR integration",            "✓", "Partial", "✓", "✓", "Manual"),
    ("Executable playbooks",        "✓", "Partial", "✓", "✓", "✗"),
    ("Threat intelligence",         "✓", "Partial", "✓", "✓", "Partial"),
    ("Unified platform",            "✓", "✗", "✗", "✓", "✗"),
    ("Open-source / free",          "✓", "✓", "✓", "✗", "✓"),
    ("NIST 800-61 lifecycle",       "✓", "✗", "Partial", "✓", "✗"),
]
for i, row_data in enumerate(comp_data):
    row = comp_table.rows[i + 1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        if i % 2 == 0:
            set_cell_bg(row.cells[j], "F0F9FF")
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                if val == "✓":
                    run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
                elif val == "✗":
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

doc.add_paragraph()

add_page_break()

# ══════════════════════════════════════════════════════════
# 6. CONCLUSION AND FUTURE WORK
# ══════════════════════════════════════════════════════════
add_heading("6. Conclusion and Future Work", 1)

add_heading("6.1 Conclusion", 2)
add_body(
    "SentriX demonstrates that a fully integrated, AI-driven Security Operations Center platform can "
    "be built using open-source tools without commercial licensing costs. The project successfully "
    "addressed the core problem of SOC fragmentation and alert fatigue by delivering a unified "
    "environment that covers the complete security operations workflow — from real-time alert ingestion "
    "and automated enrichment, through AI-assisted analysis and structured incident management, to "
    "executable response playbooks capable of performing real containment actions."
)
add_body(
    "The platform's integration of Anthropic Claude as an embedded AI analyst represents a meaningful "
    "advancement over conventional SOC tooling. Analysts at all role levels can interrogate the system "
    "in natural language, receive structured incident summaries, and obtain actionable recommendations "
    "grounded in both live SOC data and a curated security knowledge base. This directly reduces "
    "cognitive burden on analysts and shortens the time from detection to informed response."
)
add_body(
    "The implementation of the NIST SP 800-61 lifecycle within an executable, structured playbook "
    "system converts what is traditionally a documentation-driven process into an operational "
    "capability. The role-based access model ensures that junior analysts (L1) can triage and "
    "escalate alerts effectively while senior analysts (L2), incident responders, and administrators "
    "retain control over sensitive operations including playbook execution, rule configuration, "
    "and system integration management."
)
add_body(
    "SentriX achieves its primary objective: a deployable, feature-complete SOC platform that lowers "
    "the barrier to professional-grade security operations for organizations that cannot access "
    "commercial alternatives — validated through end-to-end testing across real endpoint monitoring, "
    "automated triage, AI-assisted analysis, and live playbook execution."
)

add_heading("6.2 Future Work", 2)
add_bullet("Scalable Database Backend: ", "Migrating from SQLite to PostgreSQL would enable concurrent write access, connection pooling, and horizontal scaling — prerequisites for production deployment in enterprise environments serving multiple simultaneous analyst sessions across large alert volumes.")
add_bullet("Advanced AI Threat Correlation: ", "Enhancing the AI engine with cross-incident correlation capabilities, allowing the model to identify attack campaigns spanning multiple incidents and endpoints — enabling detection of coordinated, multi-stage attacks that individual alert analysis would miss.")
add_bullet("Dedicated Mobile Application: ", "Developing a native mobile application for iOS and Android that installs directly on a device and self-configures its connection to the SentriX server upon first launch, providing real-time push notifications, alert review, and playbook approval for on-call responders.")
add_bullet("Digital Forensics Module: ", "Adding an automated forensic evidence collection capability that captures running processes, network connections, and recently modified files from affected endpoints when an incident is declared, preserving volatile evidence with a verified hash and chain of custody record.")
add_bullet("Threat Hunting Dashboard: ", "Implementing a proactive threat hunting interface where analysts can write and execute structured queries against historical alert data, search for indicators of compromise across the full event timeline, and save hunting queries as reusable templates.")
add_bullet("MITRE ATT&CK Heatmap: ", "Generating a live ATT&CK Navigator heatmap from ingested alerts, visually mapping which tactics and techniques have been observed in the environment and helping SOC managers prioritize defensive investments.")
add_bullet("Multi-Factor Authentication: ", "Adding TOTP-based two-factor authentication for all user roles to strengthen access security, particularly for privileged operations such as playbook execution and system configuration changes.")

add_page_break()

# ══════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════
add_heading("References", 1)

references = [
    "Al-Mhiqani, M. N., Ahmad, R., Yassin, W., Hassan, A., Abidin, Z. Z., Ali, N. S., & Yunos, Z. (2020). Cyber-security incidents: A review cases in cyber-physical systems. International Journal of Advanced Computer Science and Applications, 9(1), 499–508.",
    "Andrade, R. O., Yoo, S. G., Tello-Oquendo, L., & Ortiz-Garcés, I. (2021). A comprehensive study of ransomware attacks: A growing global threat. International Journal of Advanced Computer Science and Applications, 12(1).",
    "Anthropic. (2024). Claude API documentation. https://docs.anthropic.com",
    "Cichonski, P., Millar, T., Grance, T., & Scarfone, K. (2012). Computer security incident handling guide (NIST SP 800-61 Rev. 2). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.SP.800-61r2",
    "Fang, R., Bindu, R., Gupta, A., & Kang, D. (2024). LLM agents can autonomously exploit one-day vulnerabilities. arXiv preprint arXiv:2404.08144.",
    "Handa, A., Sharma, A., & Shukla, S. K. (2019). Machine learning in cybersecurity: A review. WIREs Data Mining and Knowledge Discovery, 9(4), e1306. https://doi.org/10.1002/widm.1306",
    "IBM Security. (2023). Cost of a data breach report 2023. IBM Corporation. https://www.ibm.com/reports/data-breach",
    "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Küttler, H., Lewis, M., Yih, W., Rocktäschel, T., Riedel, S., & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459–9474.",
    "MITRE Corporation. (2024). MITRE ATT&CK framework. https://attack.mitre.org",
    "Ponemon Institute. (2017). Improving the effectiveness of the SOC. Ponemon Institute LLC.",
    "TheHive Project. (2024). TheHive: Security incident response platform. https://thehive-project.org",
    "VirusTotal. (2024). VirusTotal API v3 documentation. https://developers.virustotal.com",
    "Wazuh, Inc. (2024). Wazuh open source security platform documentation. https://documentation.wazuh.com",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after       = Pt(6)
    run = p.add_run(ref)
    run.font.size = Pt(10)

# ── Save ──────────────────────────────────────────────────
output_path = "/workspaces/SentriX/SentriX_Final_Report.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
